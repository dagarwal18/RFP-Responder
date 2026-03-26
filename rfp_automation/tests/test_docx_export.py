from __future__ import annotations

import importlib.util
from pathlib import Path

from docx import Document


_MD_TO_DOCX_PATH = Path(__file__).resolve().parents[2] / "scripts" / "md_to_docx.py"
_MD_TO_DOCX_SPEC = importlib.util.spec_from_file_location("md_to_docx", _MD_TO_DOCX_PATH)
_MD_TO_DOCX = importlib.util.module_from_spec(_MD_TO_DOCX_SPEC)
assert _MD_TO_DOCX_SPEC and _MD_TO_DOCX_SPEC.loader
_MD_TO_DOCX_SPEC.loader.exec_module(_MD_TO_DOCX)

convert_md_to_docx = _MD_TO_DOCX.convert_md_to_docx


def test_convert_md_to_docx_generates_editable_word_document(tmp_path: Path):
    markdown = """# Proposal Body

This is a **strong** response with `inline code`.

- First bullet
- Second bullet

| ID | Requirement | Status |
| --- | --- | --- |
| TR-001 | Support SSO | Compliant |
"""
    input_path = tmp_path / "proposal.md"
    output_path = tmp_path / "proposal.docx"
    input_path.write_text(markdown, encoding="utf-8")

    convert_md_to_docx(
        input_path=str(input_path),
        output_path=str(output_path),
        rfp_title="Sample RFP",
        client_name="Example Client",
        company_name="Example Company",
        include_cover=False,
    )

    assert output_path.is_file()

    document = Document(output_path)
    paragraph_text = "\n".join(p.text for p in document.paragraphs if p.text.strip())

    assert "Proposal Body" in paragraph_text
    assert "First bullet" in paragraph_text
    assert "strong response" in paragraph_text
    assert len(document.tables) == 1
    assert document.tables[0].cell(1, 0).text == "TR-001"
    assert document.tables[0].cell(1, 2).text == "Compliant"
