#!/usr/bin/env python3
"""
MD-to-DOCX Converter for proposal exports.

Converts proposal markdown into an editable Word document with headings,
lists, tables, code blocks, and embedded images.
"""

from __future__ import annotations

import argparse
import logging
import re
import sys
from datetime import datetime
from pathlib import Path

logger = logging.getLogger(__name__)

try:
    from docx import Document
    from docx.enum.section import WD_SECTION_START
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Inches, Pt
except ImportError:
    Document = None  # type: ignore[assignment]

try:
    from md_to_pdf import _scrub_markdown
except ImportError:
    sys.path.insert(0, str(Path(__file__).resolve().parent))
    from md_to_pdf import _scrub_markdown


def _count_table_columns(line: str) -> int:
    stripped = line.strip()
    pipes = stripped.count("|")
    if pipes < 2:
        return 0
    if stripped.startswith("|") and stripped.endswith("|"):
        return max(0, pipes - 1)
    if not stripped.startswith("|") and not stripped.endswith("|"):
        return pipes + 1
    return pipes


def _is_table_separator(line: str) -> bool:
    normalized = line.strip().replace(" ", "")
    return bool(normalized) and set(normalized) <= {"|", "-", ":"}


def _split_table_row(line: str) -> list[str]:
    stripped = line.strip().strip("|")
    if not stripped:
        return []
    return [cell.strip() for cell in stripped.split("|")]


def _resolve_asset_path(asset_path: str, base_dir: Path) -> Path | None:
    candidate = asset_path.strip().strip("<>").strip().strip('"').strip("'")
    if not candidate or candidate.startswith(("http://", "https://", "data:")):
        return None

    normalized = Path(candidate)
    if normalized.exists():
        return normalized

    resolved = (base_dir / candidate).resolve()
    if resolved.exists():
        return resolved
    return None


def _add_page_number(paragraph) -> None:
    run = paragraph.add_run()
    fld_char_begin = OxmlElement("w:fldChar")
    fld_char_begin.set(qn("w:fldCharType"), "begin")

    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = "PAGE"

    fld_char_end = OxmlElement("w:fldChar")
    fld_char_end.set(qn("w:fldCharType"), "end")

    run._r.append(fld_char_begin)
    run._r.append(instr_text)
    run._r.append(fld_char_end)


def _apply_default_styles(document: Document) -> None:
    normal = document.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(11)

    for style_name in ("List Bullet", "List Number", "Quote"):
        if style_name in document.styles:
            document.styles[style_name].font.name = "Times New Roman"
            document.styles[style_name].font.size = Pt(11)


def _add_cover_page(
    document: Document,
    rfp_title: str,
    client_name: str,
    company_name: str,
) -> None:
    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(18)
    run = title.add_run(company_name)
    run.bold = True
    run.font.size = Pt(22)

    heading = document.add_paragraph()
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    heading.paragraph_format.space_after = Pt(14)
    run = heading.add_run(rfp_title)
    run.bold = True
    run.font.size = Pt(20)

    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(24)
    run = subtitle.add_run("Technical & Commercial Proposal")
    run.italic = True
    run.font.size = Pt(13)

    meta = document.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.add_run("Submitted to: ").bold = True
    meta.add_run(client_name)
    meta.add_run("\nSubmitted by: ").bold = True
    meta.add_run(company_name)
    meta.add_run("\nDate: ").bold = True
    meta.add_run(datetime.now().strftime("%B %d, %Y"))

    document.add_section(WD_SECTION_START.NEW_PAGE)


def _set_header_footer(section, rfp_title: str, company_name: str) -> None:
    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    header.text = f"{company_name} - {rfp_title}"

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.add_run("Page ")
    _add_page_number(footer)


def _add_inline_runs(paragraph, text: str) -> None:
    token_re = re.compile(
        r"(`[^`]+`|\*\*\*[^*]+\*\*\*|\*\*[^*]+\*\*|__[^_]+__|\[[^\]]+\]\([^)]+\)|(?<!\*)\*[^*\n]+\*(?!\*)|(?<!_)_[^_\n]+_(?!_))"
    )

    position = 0
    for match in token_re.finditer(text):
        if match.start() > position:
            paragraph.add_run(text[position:match.start()])

        token = match.group(0)
        if token.startswith("`") and token.endswith("`"):
            run = paragraph.add_run(token[1:-1])
            run.font.name = "Courier New"
            run.font.size = Pt(10)
        elif token.startswith("***") and token.endswith("***"):
            run = paragraph.add_run(token[3:-3])
            run.bold = True
            run.italic = True
        elif (token.startswith("**") and token.endswith("**")) or (
            token.startswith("__") and token.endswith("__")
        ):
            run = paragraph.add_run(token[2:-2])
            run.bold = True
        elif token.startswith("[") and "](" in token and token.endswith(")"):
            label, url = token[1:-1].split("](", 1)
            run = paragraph.add_run(label)
            run.underline = True
            paragraph.add_run(f" ({url})")
        else:
            run = paragraph.add_run(token[1:-1])
            run.italic = True
        position = match.end()

    if position < len(text):
        paragraph.add_run(text[position:])


def _add_paragraph(document: Document, text: str, style: str | None = None, indent_level: int = 0) -> None:
    paragraph = document.add_paragraph(style=style)
    if indent_level:
        paragraph.paragraph_format.left_indent = Inches(0.25 * indent_level)
    _add_inline_runs(paragraph, text)


def _add_code_block(document: Document, code_lines: list[str]) -> None:
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(6)
    paragraph.paragraph_format.space_after = Pt(6)
    run = paragraph.add_run("\n".join(code_lines))
    run.font.name = "Courier New"
    run.font.size = Pt(9)


def _add_image(document: Document, source: str, base_dir: Path) -> bool:
    resolved = _resolve_asset_path(source, base_dir)
    if not resolved:
        return False

    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    try:
        paragraph.add_run().add_picture(str(resolved), width=Inches(6.25))
    except Exception as exc:
        logger.warning("Failed to embed image %s: %s", resolved, exc)
        return False
    return True


def _add_table(document: Document, table_lines: list[str]) -> None:
    rows = [_split_table_row(line) for line in table_lines if line.strip()]
    if len(rows) < 2:
        return

    header = rows[0]
    body = [row for row in rows[1:] if row and not _is_table_separator("| " + " | ".join(row) + " |")]
    if not header or not body:
        return

    col_count = len(header)
    table = document.add_table(rows=1 + len(body), cols=col_count)
    table.style = "Table Grid"

    for idx, cell_text in enumerate(header):
        cell = table.rows[0].cells[idx]
        cell.text = cell_text
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True

    for row_index, row_values in enumerate(body, start=1):
        row = table.rows[row_index]
        padded = row_values[:col_count] + [""] * max(0, col_count - len(row_values))
        for col_index, cell_text in enumerate(padded[:col_count]):
            row.cells[col_index].text = cell_text


def convert_md_to_docx(
    input_path: str,
    output_path: str,
    rfp_title: str = "RFP Response Proposal",
    client_name: str = "Client",
    company_name: str = "Proposing Company",
    include_cover: bool = True,
) -> str:
    if Document is None:
        print(
            "ERROR: Missing dependency 'python-docx'.\n"
            "Install with: pip install python-docx",
            file=sys.stderr,
        )
        sys.exit(1)

    md_path = Path(input_path)
    if not md_path.exists():
        print(f"ERROR: Input file not found: {input_path}", file=sys.stderr)
        sys.exit(1)

    md_text = _scrub_markdown(md_path.read_text(encoding="utf-8"))
    lines = md_text.splitlines()

    document = Document()
    _apply_default_styles(document)

    section = document.sections[0]
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)

    if include_cover:
        document.sections[0].different_first_page_header_footer = True
        _add_cover_page(document, rfp_title, client_name, company_name)

    for doc_section in document.sections:
        _set_header_footer(doc_section, rfp_title, company_name)

    paragraph_buffer: list[str] = []
    code_buffer: list[str] = []
    table_buffer: list[str] = []
    in_code_block = False

    def flush_paragraph() -> None:
        nonlocal paragraph_buffer
        if paragraph_buffer:
            text = " ".join(part.strip() for part in paragraph_buffer if part.strip()).strip()
            if text:
                _add_paragraph(document, text)
            paragraph_buffer = []

    def flush_table() -> None:
        nonlocal table_buffer
        if table_buffer:
            _add_table(document, table_buffer)
            table_buffer = []

    for raw_line in lines:
        line = raw_line.rstrip()
        stripped = line.strip()

        if in_code_block:
            if stripped.startswith("```"):
                _add_code_block(document, code_buffer)
                code_buffer = []
                in_code_block = False
            else:
                code_buffer.append(line)
            continue

        if stripped.startswith("```"):
            flush_paragraph()
            flush_table()
            in_code_block = True
            code_buffer = []
            continue

        if _count_table_columns(line) >= 2:
            flush_paragraph()
            table_buffer.append(line)
            continue

        flush_table()

        if not stripped:
            flush_paragraph()
            continue

        if stripped.startswith("<div ") or stripped == "</div>":
            flush_paragraph()
            continue

        html_image_match = re.search(r'<img [^>]*src="([^"]+)"', stripped)
        md_image_match = re.fullmatch(r"!\[[^\]]*\]\((.+?)\)", stripped)
        if html_image_match or md_image_match:
            flush_paragraph()
            source = html_image_match.group(1) if html_image_match else md_image_match.group(1)
            if not _add_image(document, source, md_path.parent):
                _add_paragraph(document, f"[Image not available: {source}]")
            continue

        heading_match = re.match(r"^(#{1,6})\s+(.*)$", stripped)
        if heading_match:
            flush_paragraph()
            level = min(len(heading_match.group(1)), 4)
            _add_paragraph(document, heading_match.group(2).strip(), style=f"Heading {level}")
            continue

        quote_match = re.match(r"^>\s?(.*)$", stripped)
        if quote_match:
            flush_paragraph()
            _add_paragraph(document, quote_match.group(1).strip(), style="Quote")
            continue

        list_match = re.match(r"^(\s*)([-*+]|\d+\.)\s+(.*)$", line)
        if list_match:
            flush_paragraph()
            indent = max(0, len(list_match.group(1).replace("\t", "    ")) // 2)
            marker = list_match.group(2)
            style = "List Number" if marker.endswith(".") and marker[:-1].isdigit() else "List Bullet"
            _add_paragraph(document, list_match.group(3).strip(), style=style, indent_level=indent)
            continue

        if stripped in {"---", "***", "___"}:
            flush_paragraph()
            document.add_paragraph("_" * 30)
            continue

        paragraph_buffer.append(line)

    flush_paragraph()
    flush_table()
    if in_code_block and code_buffer:
        _add_code_block(document, code_buffer)

    out_path = Path(output_path)
    out_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(out_path)
    file_size_kb = out_path.stat().st_size / 1024
    logger.info("Generated DOCX: %s (%.1f KB)", out_path, file_size_kb)
    print(f"SUCCESS: DOCX generated at {out_path} ({file_size_kb:.1f} KB)")
    return str(out_path.resolve())


def main() -> None:
    parser = argparse.ArgumentParser(
        description="Convert Markdown output to DOCX",
        formatter_class=argparse.RawDescriptionHelpFormatter,
    )
    parser.add_argument("input", help="Path to the input Markdown file")
    parser.add_argument("output", help="Path for the output DOCX file")
    parser.add_argument(
        "--rfp-title",
        default="RFP Response Proposal",
        help="RFP title for the cover page",
    )
    parser.add_argument(
        "--client-name",
        default="Client",
        help="Client/issuing organization name",
    )
    parser.add_argument(
        "--company-name",
        default="Proposing Company",
        help="Proposing company name",
    )
    parser.add_argument(
        "--no-cover",
        action="store_true",
        help="Skip the cover page",
    )
    parser.add_argument(
        "--verbose",
        "-v",
        action="store_true",
        help="Enable verbose logging",
    )

    args = parser.parse_args()

    logging.basicConfig(
        level=logging.DEBUG if args.verbose else logging.INFO,
        format="%(levelname)s | %(message)s",
    )

    convert_md_to_docx(
        input_path=args.input,
        output_path=args.output,
        rfp_title=args.rfp_title,
        client_name=args.client_name,
        company_name=args.company_name,
        include_cover=not args.no_cover,
    )


if __name__ == "__main__":
    main()
